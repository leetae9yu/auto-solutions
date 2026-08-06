# /// script
# requires-python = ">=3.11"
# dependencies = ["PyMuPDF>=1.26", "python-docx>=1.2"]
# ///
# ─── How to run ───
# uv run scripts/pdf_to_docx.py input.pdf output.docx
"""Convert a PDF into a Word-compatible image-backed DOCX package."""

from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Mm, Pt

RENDER_DPI = 144
EXPECTED_ARGUMENT_COUNT = 3
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
IMAGE_WIDTH_MM = 208.5
IMAGE_HEIGHT_MM = 295.5


class ConversionError(RuntimeError):
    """Raised when the input cannot produce a document."""

    @classmethod
    def empty_pdf(cls) -> ConversionError:
        """Create an error for a PDF without pages."""
        return cls("PDF has no pages")

    @classmethod
    def invalid_input(cls) -> ConversionError:
        """Create an error for a non-PDF input."""
        return cls("input must be a PDF")


def _render_pages(source: Path) -> tuple[bytes, ...]:
    scale = RENDER_DPI / 72
    with pymupdf.open(source) as document:
        if document.page_count == 0:
            raise ConversionError.empty_pdf()
        return tuple(
            page.get_pixmap(matrix=pymupdf.Matrix(scale, scale), alpha=False).tobytes("png")
            for page in document
        )


def build_docx(source: Path, output: Path) -> None:
    """Render every PDF page and place it on a matching A4 DOCX page."""
    if source.suffix.lower() != ".pdf":
        raise ConversionError.invalid_input()
    pages = _render_pages(source)
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)
    document.core_properties.title = source.stem

    for index, image in enumerate(pages):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run()
        _ = run.add_picture(
            BytesIO(image),
            width=Mm(IMAGE_WIDTH_MM),
            height=Mm(IMAGE_HEIGHT_MM),
        )
        if index < len(pages) - 1:
            run.add_break(WD_BREAK.PAGE)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def main() -> int:
    """Run the PDF-to-DOCX command."""
    if len(sys.argv) != EXPECTED_ARGUMENT_COUNT:
        _ = sys.stderr.write("usage: pdf_to_docx.py INPUT.pdf OUTPUT.docx\n")
        return 2
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    try:
        build_docx(source, output)
    except (pymupdf.FileDataError, FileNotFoundError, ConversionError) as error:
        _ = sys.stderr.write(f"{error}\n")
        return 1
    _ = sys.stdout.write(f"created {output}\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
